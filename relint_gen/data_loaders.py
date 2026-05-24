"""Loaders para cada fonte de dado do CompStat.

Cada função retorna um DataFrame (ou GeoDataFrame) pandas/geopandas, com geometria
em WGS84 (EPSG:4326) quando aplicável. Pequenos toques de limpeza são feitos aqui
(encoding, separadores, decimais), mas a lógica de cruzamento fica no filters.py.
"""

from __future__ import annotations

from functools import lru_cache
from pathlib import Path

import geopandas as gpd
import pandas as pd
from shapely import wkt
from shapely.geometry import Point

from . import config


# --- ocorrências (furto/roubo georreferenciadas) ----------------------------

@lru_cache(maxsize=1)
def load_ocorrencias() -> gpd.GeoDataFrame:
    """Registros de ocorrências criminais (furto/roubo)."""
    df = pd.read_csv(
        config.OCORRENCIAS_CSV,
        dtype={"id_criptografado": "string"},
    )
    # Algumas linhas têm lat/long como NaN ou 0 — filtra.
    df = df[df["longitude"].notna() & df["latitude"].notna()].copy()
    df = df[(df["longitude"] != 0) & (df["latitude"] != 0)]
    geom = gpd.points_from_xy(df["longitude"], df["latitude"], crs=config.CRS_WGS84)
    gdf = gpd.GeoDataFrame(df, geometry=geom, crs=config.CRS_WGS84)
    # Tipos úteis
    for col in ("ano", "mes", "hora", "aisp", "risp", "delito"):
        if col in gdf.columns:
            gdf[col] = pd.to_numeric(gdf[col], errors="coerce").astype("Int64")
    return gdf


# --- disque denúncia (qualitativo) ------------------------------------------

@lru_cache(maxsize=1)
def load_denuncias() -> gpd.GeoDataFrame:
    """Denúncias do Disque Denúncia.

    Arquivo com separador ; , encoding cp1252 (chars latinos cobertos), e decimais
    com vírgula nas colunas latitude/longitude.
    """
    df = pd.read_csv(
        config.DENUNCIAS_CSV,
        sep=";",
        encoding="cp1252",
        dtype="string",
        low_memory=False,
    )
    # `dtype=string` faz o `decimal=','` virar no-op; convertemos manualmente.
    for col in ("latitude", "longitude"):
        df[col] = (
            df[col].str.replace(",", ".", regex=False)
                  .pipe(pd.to_numeric, errors="coerce")
        )
    df["data_denuncia"] = pd.to_datetime(df["data_denuncia"], errors="coerce")

    has_geo = df["latitude"].notna() & df["longitude"].notna()
    df_geo = df[has_geo].copy()
    df_geo["geometry"] = gpd.points_from_xy(
        df_geo["longitude"], df_geo["latitude"], crs=config.CRS_WGS84
    )
    return gpd.GeoDataFrame(df_geo, geometry="geometry", crs=config.CRS_WGS84)


# --- fatores urbanos --------------------------------------------------------

@lru_cache(maxsize=1)
def load_fatores_urbanos() -> gpd.GeoDataFrame:
    """Fatores ambientais/urbanos mapeados em campo."""
    df = pd.read_csv(config.FATORES_CSV, low_memory=False)
    # As coordenadas estão como coordenada_x/y (long/lat).
    df = df[df["coordenada_x"].notna() & df["coordenada_y"].notna()].copy()
    df["longitude"] = pd.to_numeric(df["coordenada_y"], errors="coerce")
    df["latitude"] = pd.to_numeric(df["coordenada_x"], errors="coerce")
    df = df.dropna(subset=["longitude", "latitude"])
    geom = gpd.points_from_xy(df["longitude"], df["latitude"], crs=config.CRS_WGS84)
    return gpd.GeoDataFrame(df, geometry=geom, crs=config.CRS_WGS84)


# --- câmeras CIVITAS/COR nas áreas FM ---------------------------------------

@lru_cache(maxsize=1)
def load_cameras() -> gpd.GeoDataFrame:
    """Câmeras instaladas nas áreas da Força Municipal."""
    df = pd.read_csv(config.CAMERAS_CSV)
    df["geometry"] = df["geometry"].map(wkt.loads)
    return gpd.GeoDataFrame(df, geometry="geometry", crs=config.CRS_WGS84)


# --- polígonos das áreas FM (shapefile) -------------------------------------

@lru_cache(maxsize=1)
def load_areas_fm() -> gpd.GeoDataFrame:
    """Polígonos das 8 áreas da Força Municipal."""
    gdf = gpd.read_file(config.SHAPE_FM)
    # Garante CRS conhecido.
    if gdf.crs is None:
        gdf = gdf.set_crs(config.CRS_WGS84)
    else:
        gdf = gdf.to_crs(config.CRS_WGS84)
    # Renomeia campo `nome_subar` (truncado pelo formato DBF) para algo legível.
    if "nome_subar" in gdf.columns:
        gdf = gdf.rename(columns={"nome_subar": "nome_area"})
    return gdf


# --- domínio territorial (facções) ------------------------------------------

@lru_cache(maxsize=1)
def load_dominio_territorial() -> gpd.GeoDataFrame:
    """Polígonos de domínio de organizações criminosas."""
    df = pd.read_csv(config.DOMINIO_CSV)
    df["geometry"] = df["geometria"].map(wkt.loads)
    return gpd.GeoDataFrame(
        df.drop(columns=["geometria"]), geometry="geometry", crs=config.CRS_WGS84
    )


# --- censo pessoas em situação de rua (PSR) ---------------------------------

@lru_cache(maxsize=1)
def load_cpsr() -> dict[int, pd.DataFrame]:
    """Censo PSR — devolve uma planilha por ano (2020, 2022, 2024)."""
    sheets = pd.read_excel(config.CPSR_XLSX, sheet_name=None)
    # As abas costumam ser nomeadas por ano; normaliza.
    out: dict[int, pd.DataFrame] = {}
    for name, frame in sheets.items():
        for token in name.split():
            if token.isdigit() and len(token) == 4:
                out[int(token)] = frame
                break
    return out or {0: pd.concat(sheets.values(), ignore_index=True)}


# --- RELINTs anteriores (texto) ---------------------------------------------

def load_relints_textos() -> dict[str, str]:
    """Lê o texto puro dos RELINTs existentes (.docx)."""
    from docx import Document  # import local — docx é pesado

    out: dict[str, str] = {}
    for path in sorted(Path(config.RELINTS_DIR).glob("*.docx")):
        doc = Document(str(path))
        texto = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        out[path.name] = texto
    return out


def load_relint_referencia(area_nome: str) -> str | None:
    """Texto do RELINT de referência para uma área (usado como style guide)."""
    from docx import Document

    nome_arquivo = config.RELINT_REFERENCIA.get(area_nome)
    if nome_arquivo is None:
        return None
    path = config.RELINTS_DIR / nome_arquivo
    if not path.exists():
        return None
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


# --- dicionário de dados ----------------------------------------------------

@lru_cache(maxsize=1)
def load_dicionario() -> dict[str, pd.DataFrame]:
    """Dicionário de dados — uma aba por dataset."""
    return pd.read_excel(config.DICIONARIO_XLSX, sheet_name=None)
