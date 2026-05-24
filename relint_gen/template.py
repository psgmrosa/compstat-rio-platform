"""Estrutura e renderizador do Relatório Analítico de Área do CompStat.

O target é o documento mostrado no Anexo do briefing técnico (págs. 11-16,
"Rua Lauro Müller – Av. Gen. Severiano – Av. Venceslau Brás"). Difere dos
RELINTs em `relints/`, que são INPUTS qualitativos (Relatórios de Inteligência
da FM) e alimentam a seção "Dinâmica Criminal" deste relatório.

Estrutura:
    Cabeçalho + Identificação da Área (AISP, DP, BPM, ORCRIM…)
    Resumo Executivo (tabela com 4 perguntas norteadoras)
    1. Ocorrências Criminais (indicadores + ranking + tipos + análise temporal)
    2. Dinâmica Criminal (síntese qualitativa via Claude sobre DD + RELINT)
    3. Efetivo Empregado — Força Municipal (situação atual + sugestão)
    4. Fatores de Incidência Criminal (tabela fator → responsável) + câmeras
    5. Plano de Ação e Responsabilização (ações + responsável + prazo + status)
"""

from __future__ import annotations

from dataclasses import asdict, dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement


# --- dataclasses do conteúdo ------------------------------------------------

@dataclass
class IdentificacaoArea:
    area_fm: str
    aisp: str = ""
    bairros: str = ""           # "Botafogo/Urca/Copacabana"
    dp: str = ""                # "10ª/12ª"
    bpm: str = ""               # "2º/19º"
    base_fm: str = ""           # "Litorânea"
    subprefeitura: str = ""     # "Zona Sul"
    num_trechos_criticos: int = 0
    influencia_orcrim: str = ""


@dataclass
class IndicadoresPeriodo:
    periodo: str                # "01/01/2023 a 31/12/2024"
    roubos: int
    furtos: int
    total: int
    ranking_areas_fm: str       # "7º lugar (5,1%)"
    variacao_periodo_anterior: str = "N/A"


@dataclass
class DistribuicaoTipo:
    ranking: int
    tipo: str                   # "Roubo a transeunte"
    quantidade: int
    pct: float = 0.0


@dataclass
class AnaliseTemporal:
    """Heatmap 7×24 (dia da semana × hora) + sumário textual."""
    heatmap: list[list[int]]    # 7 linhas (dom..sab) × 24 colunas (00..23)
    periodo_predominante: str   # "Tarde-noite (16h-22h), com pico às 19h"
    dia_horario_critico: str    # "Sexta-feira, 19h-20h"
    descricao: str = ""


@dataclass
class DinamicaCriminal:
    """Síntese gerada pelo Claude a partir do Disque Denúncia + RELINTs."""
    sintese: str
    modalidade: list[str] = field(default_factory=list)   # ["a pé", "motocicleta"]
    rotas_fuga: list[str] = field(default_factory=list)
    pontos_receptacao: list[str] = field(default_factory=list)
    perfil_suspeitos: str = ""
    fonte: str = "Disque Denúncia + RELINT da FM"


@dataclass
class EfetivoFM:
    agentes_por_turno: str = ""
    locais_cobertura: str = ""
    horario_cobertura: str = ""
    dias_cobertura: str = ""
    modalidade_emprego: str = ""    # "a pé", "moto", "viatura", "mista"
    sugestao_agentes: str = ""
    sugestao_locais: str = ""
    sugestao_horario: str = ""
    sugestao_dias: str = ""
    sugestao_modalidade: str = ""
    justificativa: str = ""


@dataclass
class FatorIncidencia:
    fator: str                  # "Iluminação Pública"
    descricao: str
    responsavel: str            # "RioLuz" / "Comlurb" / etc.


@dataclass
class AcaoPlano:
    acao: str
    responsavel: str
    prazo: str = ""
    status: str = "Pendente"


@dataclass
class PerguntaNorteadora:
    pergunta: str
    diagnostico: str            # com base nos dados
    operacao_sugerida: str
    observacao: str = ""


@dataclass
class ResumoExecutivo:
    perguntas: list[PerguntaNorteadora] = field(default_factory=list)


@dataclass
class RelatorioAnalitico:
    identificacao: IdentificacaoArea
    indicadores: IndicadoresPeriodo
    distribuicao: list[DistribuicaoTipo]
    temporal: AnaliseTemporal
    dinamica: DinamicaCriminal
    efetivo: EfetivoFM
    fatores: list[FatorIncidencia]
    plano_acao: list[AcaoPlano]
    resumo_executivo: ResumoExecutivo
    cameras_total: int = 0
    cameras_descricao: str = ""
    heatmap_image: str | None = None   # path para PNG do mapa
    metadados: dict = field(default_factory=dict)

    def to_dict(self) -> dict:
        return asdict(self)


# --- helpers de estilo do .docx ---------------------------------------------

DIAS_SEMANA = ["Domingo", "Segunda", "Terça", "Quarta", "Quinta", "Sexta", "Sábado"]
COR_HEADER = RGBColor(0x1F, 0x3A, 0x5F)        # azul escuro
COR_TEXTO_BRANCO = RGBColor(0xFF, 0xFF, 0xFF)


def _shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _heatmap_cell_color(value: int, vmax: int) -> str:
    """Cor do heatmap (branco → vermelho)."""
    if vmax <= 0 or value <= 0:
        return "FFFFFF"
    ratio = min(1.0, value / vmax)
    # branco (#FFFFFF) → vermelho saturado (#B22222)
    r = 255
    g = int(255 - (255 - 0x22) * ratio)
    b = int(255 - (255 - 0x22) * ratio)
    return f"{r:02X}{g:02X}{b:02X}"


def _add_section_heading(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = COR_HEADER


def _add_subheading(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)


def _set_cell_text(cell, text: str, *, bold: bool = False, white: bool = False,
                   center: bool = False, size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = COR_TEXTO_BRANCO
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _two_col_table(doc, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = True
    for i, (label, value) in enumerate(rows):
        _set_cell_text(table.rows[i].cells[0], label, bold=True, white=True, size=10)
        _shade_cell(table.rows[i].cells[0], "1F3A5F")
        _set_cell_text(table.rows[i].cells[1], value, size=10)


# --- renderer ---------------------------------------------------------------

def render_relatorio(rel: RelatorioAnalitico, output_path: Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _render_cabecalho(doc, rel)
    _render_resumo_executivo(doc, rel)
    _render_secao_1_ocorrencias(doc, rel)
    _render_secao_2_dinamica(doc, rel)
    _render_secao_3_efetivo(doc, rel)
    _render_secao_4_fatores(doc, rel)
    _render_secao_5_plano(doc, rel)

    doc.save(str(output_path))
    return output_path


def _render_cabecalho(doc, rel: RelatorioAnalitico) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RELATÓRIO ANALÍTICO DE ÁREA")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = COR_HEADER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("Subsídio para Reunião de CompStat")
    s.italic = True
    s.font.size = Pt(11)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    _set_cell_text(table.rows[0].cells[0], "Área de análise",
                   bold=True, white=True, center=True)
    _shade_cell(table.rows[0].cells[0], "1F3A5F")
    _set_cell_text(table.rows[0].cells[1], rel.identificacao.area_fm, bold=True)

    if rel.heatmap_image:
        try:
            doc.add_paragraph()
            doc.add_picture(rel.heatmap_image, width=Cm(15))
        except Exception:
            pass

    doc.add_paragraph()


def _render_resumo_executivo(doc, rel: RelatorioAnalitico) -> None:
    _add_section_heading(doc, "RESUMO EXECUTIVO")
    doc.add_paragraph(
        "Síntese gerada automaticamente pela plataforma, respondendo às "
        "perguntas norteadoras da reunião de CompStat. Os diagnósticos são "
        "ancorados nos dados de ocorrências, Disque Denúncia, RELINTs e "
        "fatores urbanos georreferenciados."
    )

    perguntas = rel.resumo_executivo.perguntas or []
    if not perguntas:
        return

    table = doc.add_table(rows=1 + len(perguntas), cols=4)
    headers = ["Pergunta norteadora", "Diagnóstico (dados)",
               "Operação FM / órgãos", "Observação"]
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, white=True, center=True)
        _shade_cell(table.rows[0].cells[i], "1F3A5F")

    for i, q in enumerate(perguntas, start=1):
        _set_cell_text(table.rows[i].cells[0], q.pergunta, bold=True, size=10)
        _set_cell_text(table.rows[i].cells[1], q.diagnostico, size=10)
        _set_cell_text(table.rows[i].cells[2], q.operacao_sugerida, size=10)
        _set_cell_text(table.rows[i].cells[3], q.observacao, size=10)

    doc.add_paragraph()


def _render_secao_1_ocorrencias(doc, rel: RelatorioAnalitico) -> None:
    _add_section_heading(doc, "1. OCORRÊNCIAS CRIMINAIS")

    _add_subheading(doc, "Identificação da área")
    ident = rel.identificacao
    _two_col_table(doc, [
        ("Área FM", ident.area_fm),
        ("AISP", ident.aisp),
        ("Bairros", ident.bairros),
        ("DP", ident.dp),
        ("BPM", ident.bpm),
        ("Base FM", ident.base_fm),
        ("Subprefeitura", ident.subprefeitura),
        ("Nº de trechos críticos", str(ident.num_trechos_criticos)),
        ("Área sob influência de ORCRIM", ident.influencia_orcrim or "—"),
    ])
    doc.add_paragraph()

    _add_subheading(doc, "Indicadores do período")
    ind = rel.indicadores
    t = doc.add_table(rows=2, cols=6)
    headers = ["Período", "Roubos", "Furtos", "Total", "Ranking entre áreas FM",
               "Variação s/ período anterior"]
    for i, h in enumerate(headers):
        _set_cell_text(t.rows[0].cells[i], h, bold=True, white=True, center=True, size=9)
        _shade_cell(t.rows[0].cells[i], "1F3A5F")
    valores = [ind.periodo, f"{ind.roubos:,}".replace(",", "."),
               f"{ind.furtos:,}".replace(",", "."),
               f"{ind.total:,}".replace(",", "."),
               ind.ranking_areas_fm, ind.variacao_periodo_anterior]
    for i, v in enumerate(valores):
        _set_cell_text(t.rows[1].cells[i], v, center=True, size=10)
    doc.add_paragraph()

    if rel.distribuicao:
        _add_subheading(doc, "Distribuição de ocorrências por tipo")
        t2 = doc.add_table(rows=1 + len(rel.distribuicao), cols=4)
        for i, h in enumerate(["Ranking", "Tipo de ocorrência", "Qtd", "% do total"]):
            _set_cell_text(t2.rows[0].cells[i], h, bold=True, white=True, center=True)
            _shade_cell(t2.rows[0].cells[i], "1F3A5F")
        for i, d in enumerate(rel.distribuicao, start=1):
            _set_cell_text(t2.rows[i].cells[0], f"{d.ranking}º", center=True)
            _set_cell_text(t2.rows[i].cells[1], d.tipo)
            _set_cell_text(t2.rows[i].cells[2], f"{d.quantidade:,}".replace(",", "."),
                           center=True)
            _set_cell_text(t2.rows[i].cells[3], f"{d.pct:.1f}%", center=True)
        doc.add_paragraph()

    _add_subheading(doc, "Análise temporal")
    tt = doc.add_table(rows=8, cols=25)
    # cabeçalho de horas
    _set_cell_text(tt.rows[0].cells[0], "Dia", bold=True, white=True, center=True, size=8)
    _shade_cell(tt.rows[0].cells[0], "1F3A5F")
    for h in range(24):
        _set_cell_text(tt.rows[0].cells[h + 1], f"{h:02d}",
                       bold=True, white=True, center=True, size=8)
        _shade_cell(tt.rows[0].cells[h + 1], "1F3A5F")
    vmax = max((v for row in rel.temporal.heatmap for v in row), default=0)
    for d in range(7):
        _set_cell_text(tt.rows[d + 1].cells[0], DIAS_SEMANA[d],
                       bold=True, white=True, size=8)
        _shade_cell(tt.rows[d + 1].cells[0], "1F3A5F")
        for h in range(24):
            val = rel.temporal.heatmap[d][h] if d < len(rel.temporal.heatmap) else 0
            _set_cell_text(tt.rows[d + 1].cells[h + 1],
                           str(val) if val else "·", center=True, size=8)
            _shade_cell(tt.rows[d + 1].cells[h + 1],
                        _heatmap_cell_color(val, vmax))
    doc.add_paragraph()

    _two_col_table(doc, [
        ("Período predominante", rel.temporal.periodo_predominante),
        ("Dia / horário crítico", rel.temporal.dia_horario_critico),
        ("Descrição", rel.temporal.descricao or ""),
    ])
    doc.add_paragraph()


def _render_secao_2_dinamica(doc, rel: RelatorioAnalitico) -> None:
    _add_section_heading(doc, "2. DINÂMICA CRIMINAL")
    italic = doc.add_paragraph()
    r = italic.add_run(f"Fonte: {rel.dinamica.fonte}")
    r.italic = True
    r.font.size = Pt(9)

    p = doc.add_paragraph(rel.dinamica.sintese)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    extras: list[tuple[str, str]] = []
    if rel.dinamica.modalidade:
        extras.append(("Modalidade", ", ".join(rel.dinamica.modalidade)))
    if rel.dinamica.rotas_fuga:
        extras.append(("Rotas de fuga / escoamento", "; ".join(rel.dinamica.rotas_fuga)))
    if rel.dinamica.pontos_receptacao:
        extras.append(("Pontos de receptação", "; ".join(rel.dinamica.pontos_receptacao)))
    if rel.dinamica.perfil_suspeitos:
        extras.append(("Perfil de suspeitos", rel.dinamica.perfil_suspeitos))
    if extras:
        _two_col_table(doc, extras)
    doc.add_paragraph()


def _render_secao_3_efetivo(doc, rel: RelatorioAnalitico) -> None:
    _add_section_heading(doc, "3. EFETIVO EMPREGADO – FORÇA MUNICIPAL")
    e = rel.efetivo
    t = doc.add_table(rows=6, cols=4)
    for i, h in enumerate(["Campo", "Situação atual", "Sugestão", "Justificativa"]):
        _set_cell_text(t.rows[0].cells[i], h, bold=True, white=True, center=True)
        _shade_cell(t.rows[0].cells[i], "1F3A5F")

    linhas = [
        ("Nº de Agentes por Turno", e.agentes_por_turno, e.sugestao_agentes),
        ("Locais de Cobertura", e.locais_cobertura, e.sugestao_locais),
        ("Horário de Cobertura", e.horario_cobertura, e.sugestao_horario),
        ("Dias de Cobertura", e.dias_cobertura, e.sugestao_dias),
        ("Modalidade de Emprego", e.modalidade_emprego, e.sugestao_modalidade),
    ]
    for i, (campo, atual, sugestao) in enumerate(linhas, start=1):
        _set_cell_text(t.rows[i].cells[0], campo, bold=True, size=10)
        _set_cell_text(t.rows[i].cells[1], atual or "—", size=10)
        _set_cell_text(t.rows[i].cells[2], sugestao or "—", size=10)
        _set_cell_text(t.rows[i].cells[3], e.justificativa if i == 1 else "", size=9)
    doc.add_paragraph()


def _render_secao_4_fatores(doc, rel: RelatorioAnalitico) -> None:
    _add_section_heading(doc, "4. FATORES DE INCIDÊNCIA CRIMINAL")
    if not rel.fatores:
        doc.add_paragraph("Nenhum fator urbano relevante mapeado na área.")
        return

    t = doc.add_table(rows=1 + len(rel.fatores), cols=3)
    for i, h in enumerate(["Fator identificado", "Descrição", "Responsável"]):
        _set_cell_text(t.rows[0].cells[i], h, bold=True, white=True, center=True)
        _shade_cell(t.rows[0].cells[i], "1F3A5F")
    for i, f in enumerate(rel.fatores, start=1):
        _set_cell_text(t.rows[i].cells[0], f.fator, bold=True, size=10)
        _set_cell_text(t.rows[i].cells[1], f.descricao, size=10)
        _set_cell_text(t.rows[i].cells[2], f.responsavel, center=True, size=10)
    doc.add_paragraph()

    _add_subheading(doc, "Câmeras identificadas na área")
    _two_col_table(doc, [
        ("Quantidade", str(rel.cameras_total)),
        ("Descrição", rel.cameras_descricao or "Câmeras CIVITAS/COR mapeadas."),
    ])
    doc.add_paragraph()


def _render_secao_5_plano(doc, rel: RelatorioAnalitico) -> None:
    _add_section_heading(doc, "5. PLANO DE AÇÃO E RESPONSABILIZAÇÃO")
    italic = doc.add_paragraph()
    r = italic.add_run(
        "Pré-populado pela IA a partir das coincidências (mancha × fator × dinâmica). "
        "Revisar e formalizar na reunião de CompStat."
    )
    r.italic = True
    r.font.size = Pt(9)

    if not rel.plano_acao:
        return
    t = doc.add_table(rows=1 + len(rel.plano_acao), cols=4)
    for i, h in enumerate(["Ação acordada", "Responsável", "Prazo", "Status"]):
        _set_cell_text(t.rows[0].cells[i], h, bold=True, white=True, center=True)
        _shade_cell(t.rows[0].cells[i], "1F3A5F")
    for i, a in enumerate(rel.plano_acao, start=1):
        _set_cell_text(t.rows[i].cells[0], a.acao, size=10)
        _set_cell_text(t.rows[i].cells[1], a.responsavel, bold=True, center=True, size=10)
        _set_cell_text(t.rows[i].cells[2], a.prazo, center=True, size=10)
        _set_cell_text(t.rows[i].cells[3], a.status, center=True, size=10)


# --- schema para Claude tool_use --------------------------------------------

RELATORIO_TOOL_SCHEMA = {
    "name": "submit_relatorio_analitico",
    "description": (
        "Submete o conteúdo estruturado do Relatório Analítico de Área "
        "do CompStat. Toda afirmação textual deve ser ancorada nos dados "
        "fornecidos. Recomendações devem indicar órgão responsável."
    ),
    "input_schema": {
        "type": "object",
        "required": ["identificacao", "indicadores", "distribuicao", "temporal",
                     "dinamica", "efetivo", "fatores", "plano_acao",
                     "resumo_executivo"],
        "properties": {
            "identificacao": {
                "type": "object",
                "properties": {
                    "area_fm": {"type": "string"},
                    "aisp": {"type": "string"},
                    "bairros": {"type": "string"},
                    "dp": {"type": "string"},
                    "bpm": {"type": "string"},
                    "base_fm": {"type": "string"},
                    "subprefeitura": {"type": "string"},
                    "num_trechos_criticos": {"type": "integer"},
                    "influencia_orcrim": {"type": "string"},
                },
                "required": ["area_fm"],
            },
            "indicadores": {
                "type": "object",
                "required": ["periodo", "roubos", "furtos", "total",
                             "ranking_areas_fm"],
                "properties": {
                    "periodo": {"type": "string"},
                    "roubos": {"type": "integer"},
                    "furtos": {"type": "integer"},
                    "total": {"type": "integer"},
                    "ranking_areas_fm": {"type": "string"},
                    "variacao_periodo_anterior": {"type": "string"},
                },
            },
            "distribuicao": {
                "type": "array",
                "minItems": 1,
                "items": {
                    "type": "object",
                    "required": ["ranking", "tipo", "quantidade"],
                    "properties": {
                        "ranking": {"type": "integer"},
                        "tipo": {"type": "string"},
                        "quantidade": {"type": "integer"},
                        "pct": {"type": "number"},
                    },
                },
            },
            "temporal": {
                "type": "object",
                "required": ["periodo_predominante", "dia_horario_critico"],
                "properties": {
                    "periodo_predominante": {"type": "string"},
                    "dia_horario_critico": {"type": "string"},
                    "descricao": {"type": "string"},
                },
            },
            "dinamica": {
                "type": "object",
                "required": ["sintese"],
                "properties": {
                    "sintese": {"type": "string"},
                    "modalidade": {"type": "array", "items": {"type": "string"}},
                    "rotas_fuga": {"type": "array", "items": {"type": "string"}},
                    "pontos_receptacao": {"type": "array", "items": {"type": "string"}},
                    "perfil_suspeitos": {"type": "string"},
                },
            },
            "efetivo": {
                "type": "object",
                "properties": {
                    "agentes_por_turno": {"type": "string"},
                    "locais_cobertura": {"type": "string"},
                    "horario_cobertura": {"type": "string"},
                    "dias_cobertura": {"type": "string"},
                    "modalidade_emprego": {"type": "string"},
                    "sugestao_agentes": {"type": "string"},
                    "sugestao_locais": {"type": "string"},
                    "sugestao_horario": {"type": "string"},
                    "sugestao_dias": {"type": "string"},
                    "sugestao_modalidade": {"type": "string"},
                    "justificativa": {"type": "string"},
                },
            },
            "fatores": {
                "type": "array",
                "items": {
                    "type": "object",
                    "required": ["fator", "descricao", "responsavel"],
                    "properties": {
                        "fator": {"type": "string"},
                        "descricao": {"type": "string"},
                        "responsavel": {"type": "string"},
                    },
                },
            },
            "plano_acao": {
                "type": "array",
                "minItems": 3,
                "items": {
                    "type": "object",
                    "required": ["acao", "responsavel"],
                    "properties": {
                        "acao": {"type": "string"},
                        "responsavel": {"type": "string"},
                        "prazo": {"type": "string"},
                        "status": {"type": "string"},
                    },
                },
            },
            "resumo_executivo": {
                "type": "object",
                "required": ["perguntas"],
                "properties": {
                    "perguntas": {
                        "type": "array",
                        "minItems": 4,
                        "maxItems": 4,
                        "items": {
                            "type": "object",
                            "required": ["pergunta", "diagnostico",
                                         "operacao_sugerida"],
                            "properties": {
                                "pergunta": {"type": "string"},
                                "diagnostico": {"type": "string"},
                                "operacao_sugerida": {"type": "string"},
                                "observacao": {"type": "string"},
                            },
                        },
                    },
                },
            },
        },
    },
}


def relatorio_from_dict(d: dict, *, heatmap: list[list[int]] | None = None,
                        heatmap_image: str | None = None,
                        cameras_total: int = 0) -> RelatorioAnalitico:
    """Converte o payload do Claude num RelatorioAnalitico completo.

    Argumentos extras (heatmap, cameras_total) vêm dos analytics — o Claude
    sintetiza, mas os números vêm dos dados.
    """
    temporal = AnaliseTemporal(
        heatmap=heatmap or [[0] * 24 for _ in range(7)],
        periodo_predominante=d["temporal"]["periodo_predominante"],
        dia_horario_critico=d["temporal"]["dia_horario_critico"],
        descricao=d["temporal"].get("descricao", ""),
    )
    return RelatorioAnalitico(
        identificacao=IdentificacaoArea(**d["identificacao"]),
        indicadores=IndicadoresPeriodo(**d["indicadores"]),
        distribuicao=[DistribuicaoTipo(**x) for x in d["distribuicao"]],
        temporal=temporal,
        dinamica=DinamicaCriminal(**d["dinamica"]),
        efetivo=EfetivoFM(**d["efetivo"]),
        fatores=[FatorIncidencia(**x) for x in d["fatores"]],
        plano_acao=[AcaoPlano(**x) for x in d["plano_acao"]],
        resumo_executivo=ResumoExecutivo(
            perguntas=[PerguntaNorteadora(**q)
                       for q in d["resumo_executivo"]["perguntas"]]
        ),
        cameras_total=cameras_total,
        heatmap_image=heatmap_image,
        metadados=d.get("metadados", {}),
    )
